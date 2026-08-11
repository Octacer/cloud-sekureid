"""Text extraction endpoint (PDF native/OCR, spreadsheets, Word docs, images)."""

import os
import shutil
import subprocess
import uuid
from datetime import datetime

from fastapi import APIRouter, HTTPException, BackgroundTasks
import requests
import magic
from pdf2image import convert_from_path
from PIL import Image
import pytesseract

from src.config import EXTRACT_TEMP_DIR
from src.helpers import download_url_to_file, cleanup_file, detect_document_type
from src.models import TextExtractionRequest, TextExtractionResponse

router = APIRouter()

# Cap the number of pages we process so a huge PDF can't OCR forever.
MAX_PAGES = 50
# A page with real text yields far more than this; below it we assume the
# page is scanned/image-only and fall back to OCR.
NATIVE_CHARS_PER_PAGE_THRESHOLD = 40


@router.post("/extract-text", response_model=TextExtractionResponse)
async def extract_text(
    request_data: TextExtractionRequest,
    background_tasks: BackgroundTasks
):
    """
    Extract text from an image, PDF, spreadsheet, or Word document.

    Parameters:
    - url: Publicly accessible URL to the file

    Returns:
    - JSON with extracted text, language detected, and extraction metadata
    """
    request_id = str(uuid.uuid4())
    temp_extract_dir = EXTRACT_TEMP_DIR
    os.makedirs(temp_extract_dir, exist_ok=True)

    try:
        print(f"\n{'='*80}")
        print(f"[{request_id}] Processing text extraction request")
        print(f"{'='*80}")
        print(f"[{request_id}] URL: {request_data.url}")

        url_str = str(request_data.url)

        # Download file first
        print(f"[{request_id}] Step 1: Downloading file from URL...")
        temp_raw_file = os.path.join(temp_extract_dir, f"{request_id}_raw")
        bytes_written = download_url_to_file(url_str, temp_raw_file, timeout=30)
        file_size = os.path.getsize(temp_raw_file)
        print(f"[{request_id}] File saved: {bytes_written} bytes ({file_size} on disk)")

        # Detect file type from actual file content using magic bytes
        print(f"[{request_id}] Step 2: Detecting file type from content...")
        try:
            mime = magic.Magic(mime=True)
            detected_mime = mime.from_file(temp_raw_file)
            print(f"[{request_id}] Detected MIME type: {detected_mime}")
        except Exception as mime_error:
            print(f"[{request_id}] ERROR detecting MIME: {mime_error}")
            raise

        detected = detect_document_type(temp_raw_file, detected_mime)
        is_pdf = detected.is_pdf
        is_spreadsheet = detected.is_spreadsheet
        is_word = detected.is_word
        source_type = detected.source_type
        file_extension = detected.file_extension

        print(f"[{request_id}] Is PDF: {is_pdf}")
        print(f"[{request_id}] Is spreadsheet: {is_spreadsheet}")
        print(f"[{request_id}] Is word: {is_word}")
        print(f"[{request_id}] Source type: {source_type}")
        print(f"[{request_id}] File extension: {file_extension}")

        # Rename temporary file to proper extension
        temp_file = os.path.join(temp_extract_dir, f"{request_id}.{file_extension}")
        os.rename(temp_raw_file, temp_file)
        print(f"[{request_id}] Step 3: Renamed to {temp_file}")

        # Extract text
        extracted_text = ""
        total_pages = 0
        extraction_method = "Tesseract OCR"

        print(f"[{request_id}] Step 4: Text extraction")

        if is_pdf:
            print(f"[{request_id}] Processing PDF file")

            # Determine page count cheaply via pdfinfo (poppler) so we can cap work
            # and size the native-text threshold correctly.
            page_count = 0
            try:
                pdfinfo = subprocess.run(
                    ['pdfinfo', temp_file],
                    capture_output=True, text=True, timeout=30
                )
                for line in pdfinfo.stdout.splitlines():
                    if line.startswith('Pages:'):
                        page_count = int(line.split(':', 1)[1].strip())
                        break
            except Exception as info_error:
                print(f"[{request_id}] WARNING: pdfinfo failed: {info_error}")

            capped_pages = min(page_count, MAX_PAGES) if page_count else MAX_PAGES
            if page_count > MAX_PAGES:
                print(f"[{request_id}] PDF has {page_count} pages; capping to {MAX_PAGES}")

            # Native-first: try pdftotext -layout (preserves columns/tables).
            native_text = ""
            print(f"[{request_id}] Attempting native extraction (pdftotext -layout)...")
            try:
                pdftotext_cmd = ['pdftotext', '-layout']
                if page_count > MAX_PAGES:
                    pdftotext_cmd += ['-l', str(MAX_PAGES)]
                pdftotext_cmd += [temp_file, '-']
                result = subprocess.run(
                    pdftotext_cmd, capture_output=True, text=True, timeout=120
                )
                native_text = result.stdout or ""
                print(f"[{request_id}] Native extraction: {len(native_text.strip())} characters")
            except Exception as native_error:
                print(f"[{request_id}] WARNING: pdftotext failed: {native_error}")

            threshold = NATIVE_CHARS_PER_PAGE_THRESHOLD * max(capped_pages, 1)
            if len(native_text.strip()) >= threshold:
                # Text PDF — use the faithful native extraction.
                extracted_text = native_text
                total_pages = page_count or capped_pages
                extraction_method = "pdftotext (native)"
                print(f"[{request_id}] Using native extraction (>= threshold {threshold})")
            else:
                # Looks scanned/image-only — fall back to OCR.
                print(f"[{request_id}] Native text below threshold {threshold}; falling back to OCR")
                print(f"[{request_id}] Converting PDF to images (DPI: 200)...")
                try:
                    # Convert PDF to images (respecting the page cap)
                    images = convert_from_path(temp_file, dpi=200, last_page=MAX_PAGES)
                    total_pages = page_count or len(images)
                    print(f"[{request_id}] PDF conversion successful")
                    print(f"[{request_id}] Total pages: {total_pages}, OCR'ing {len(images)}")

                    # Extract text from each page
                    print(f"[{request_id}] Extracting text from PDF pages...")
                    page_texts = []
                    for i, image in enumerate(images, start=1):
                        print(f"[{request_id}] Processing page {i}/{len(images)}...")
                        try:
                            page_text = pytesseract.image_to_string(image)
                            text_len = len(page_text.strip())
                            print(f"[{request_id}] Page {i} extracted: {text_len} characters")
                            page_texts.append(f"--- PAGE {i} ---\n{page_text}")
                        except Exception as page_error:
                            print(f"[{request_id}] ERROR extracting page {i}: {page_error}")
                            page_texts.append(f"--- PAGE {i} ---\nERROR: {str(page_error)}")

                    extracted_text = "\n\n".join(page_texts)
                    extraction_method = "Tesseract OCR (fallback)"
                    print(f"[{request_id}] All pages processed")
                except Exception as pdf_error:
                    print(f"[{request_id}] ERROR processing PDF: {pdf_error}")
                    raise

        elif is_spreadsheet:
            print(f"[{request_id}] Processing spreadsheet file")
            try:
                import openpyxl

                # First pass: formulas (data_only=False keeps the =SUM(...) strings).
                wb = openpyxl.load_workbook(temp_file, data_only=False, read_only=True)
                # Second pass: cached computed values, if the writer saved them.
                try:
                    wb_values = openpyxl.load_workbook(temp_file, data_only=True, read_only=True)
                except Exception as values_error:
                    print(f"[{request_id}] WARNING: values pass failed: {values_error}")
                    wb_values = None

                total_pages = len(wb.sheetnames)
                print(f"[{request_id}] Workbook loaded, {total_pages} sheet(s)")

                sheet_texts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    ws_values = wb_values[sheet_name] if wb_values and sheet_name in wb_values.sheetnames else None
                    print(f"[{request_id}] Sheet '{sheet_name}': {ws.max_row} rows x {ws.max_column} cols")

                    lines = [f"--- SHEET: {sheet_name} ---"]
                    for row in ws.iter_rows():
                        for cell in row:
                            if cell.value is None:
                                continue
                            entry = f"{cell.coordinate}: {cell.value}"
                            # If the cell holds a formula, also surface its computed value.
                            if ws_values is not None and isinstance(cell.value, str) and cell.value.startswith('='):
                                computed = ws_values[cell.coordinate].value
                                if computed is not None:
                                    entry += f" => {computed}"
                            lines.append(entry)
                    sheet_texts.append("\n".join(lines))

                wb.close()
                if wb_values is not None:
                    wb_values.close()

                extracted_text = "\n\n".join(sheet_texts)
                extraction_method = "openpyxl"
                print(f"[{request_id}] Spreadsheet extraction complete")
            except Exception as xlsx_error:
                print(f"[{request_id}] ERROR processing spreadsheet: {xlsx_error}")
                raise

        elif is_word:
            print(f"[{request_id}] Processing Word document")
            try:
                import docx

                document = docx.Document(temp_file)
                parts = []

                # Body paragraphs, in document order.
                for para in document.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

                # Tables — emit each row as tab-separated cells so tabular data
                # (niches, scores, rankings) stays grader-readable.
                for t_index, table in enumerate(document.tables, start=1):
                    parts.append(f"--- TABLE {t_index} ---")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append("\t".join(cells))

                extracted_text = "\n".join(parts)
                # A .docx has no fixed page count until rendered; report 1.
                total_pages = 1
                extraction_method = "python-docx"
                print(f"[{request_id}] Word extraction complete: "
                      f"{len(document.paragraphs)} paragraphs, {len(document.tables)} table(s)")
            except Exception as docx_error:
                print(f"[{request_id}] ERROR processing Word document: {docx_error}")
                raise

        else:
            print(f"[{request_id}] Processing image file")
            try:
                # Open image and extract text
                image = Image.open(temp_file)
                print(f"[{request_id}] Image opened: size={image.size}, mode={image.mode}")

                print(f"[{request_id}] Running Tesseract OCR on image...")
                extracted_text = pytesseract.image_to_string(image)
                print(f"[{request_id}] Tesseract OCR complete")
                total_pages = 1
            except Exception as img_error:
                print(f"[{request_id}] ERROR processing image: {img_error}")
                raise

        print(f"[{request_id}] Text extraction complete: {len(extracted_text)} characters")
        if extracted_text:
            print(f"[{request_id}] Preview: {extracted_text[:200]}")
        else:
            print(f"[{request_id}] WARNING: No text extracted!")

        extracted_at = datetime.now()

        print(f"[{request_id}] Request completed successfully\n")

        # Cleanup temp file
        background_tasks.add_task(cleanup_file, temp_file, 60)

        return TextExtractionResponse(
            text=extracted_text,
            language="eng",  # Tesseract default, can be extended for language detection
            extraction_method=extraction_method,
            source_type=source_type,
            total_pages=total_pages,
            extracted_at=extracted_at.isoformat(),
            request_id=request_id
        )

    except requests.RequestException as e:
        print(f"[{request_id}] REQUEST ERROR - {type(e).__name__}: {e}")
        shutil.rmtree(temp_extract_dir, ignore_errors=True)
        raise HTTPException(
            status_code=400,
            detail=f"Failed to download file from URL: {str(e)}"
        )

    except Exception as e:
        print(f"[{request_id}] GENERAL ERROR - {type(e).__name__}: {e}")
        import traceback
        print(f"[{request_id}] Full traceback: {traceback.format_exc()}")
        shutil.rmtree(temp_extract_dir, ignore_errors=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text: {str(e)}"
        )
